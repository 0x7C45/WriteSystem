"""校验工具：字数统计、引用一致性、格式合规检查。"""

import re
from pathlib import Path
from typing import Literal

from paper_tools_mcp.utils.safe_path import safe_read, safe_path_exists, template_path_exists

# OpenXML namespace for direct XML inspection
DOCX_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _count_chinese_chars(text: str) -> int:
    """统计中文字符数（含中文标点）。"""
    return len(re.findall(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]', text))


def _count_all_chars(text: str) -> int:
    """统计所有非空白字符数。"""
    return len(re.sub(r'\s+', '', text))


def _extract_chapters(md_text: str) -> list[dict]:
    """从 Markdown 文本中按 # 标题提取章节。"""
    chapters: list[dict] = []
    current_title = "引言（无标题部分）"
    current_content: list[str] = []
    lines = md_text.split('\n')

    for line in lines:
        if line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
            if current_content:
                content = '\n'.join(current_content).strip()
                if content:
                    chapters.append({
                        "title": current_title,
                        "content": content,
                        "chinese_chars": _count_chinese_chars(content),
                        "all_chars": _count_all_chars(content),
                    })
            current_title = line.lstrip('#').strip()
            current_content = []
        else:
            current_content.append(line)

    if current_content:
        content = '\n'.join(current_content).strip()
        if content:
            chapters.append({
                "title": current_title,
                "content": content,
                "chinese_chars": _count_chinese_chars(content),
                "all_chars": _count_all_chars(content),
            })

    return chapters


def _normalize_chapter_label(label: str) -> str:
    """规范化章节标签，便于跨文件比对。"""
    return re.sub(r'\s+', '', label.strip())


def _extract_chapter_label(chapter_path: str, text: str) -> str | None:
    """从文件名或一级标题中提取“第N章”标签。"""
    candidates = [Path(chapter_path).stem]
    heading_match = re.search(r'^\s*#\s+(.+)$', text, re.MULTILINE)
    if heading_match:
        candidates.append(heading_match.group(1))

    for candidate in candidates:
        match = re.search(r'(第\s*\d+\s*章)', candidate)
        if match:
            return _normalize_chapter_label(match.group(1))
    return None


def _parse_registry_entries(reg_text: str) -> list[dict]:
    """解析引用编号注册表中的编号/cite_key/首次出现章节。"""
    entries: list[dict] = []
    seen_nums: set[int] = set()

    for line in reg_text.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r'^\|\s*-+', line):
            continue

        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if len(cells) < 2 or not cells[0].isdigit():
            continue

        num = int(cells[0])
        if num in seen_nums:
            continue
        seen_nums.add(num)
        entries.append({
            "num": num,
            "cite_key": cells[1],
            "first_chapter": cells[2] if len(cells) >= 3 else "",
        })

    return entries


def validate_word_count(
    file_path: str,
    target_total: int | None = None,
    mode: str = "chinese",
) -> dict:
    """字数统计（总量 + 按章节）。

    Args:
        file_path: Markdown 文件路径。
        target_total: 目标总字数（可选，用于对比）。
        mode: 统计模式，"chinese"（仅中文字符）或 "all"（所有非空白字符）。
    """
    if mode not in ("chinese", "all"):
        return {"error": f"无效模式: {mode}，应为 'chinese' 或 'all'"}

    try:
        path, text = safe_read(file_path)
    except ValueError as e:
        return {"error": str(e)}

    chapters = _extract_chapters(text)

    count_fn = _count_chinese_chars if mode == "chinese" else _count_all_chars
    total = count_fn(text)

    report: dict = {
        "file": file_path,
        "mode": mode,
        "total": total,
        "chapters": [
            {"title": ch["title"], "chars": ch["chinese_chars" if mode == "chinese" else "all_chars"]}
            for ch in chapters
        ],
    }

    if target_total is not None:
        diff = total - target_total
        report["target"] = target_total
        report["diff"] = diff
        report["status"] = "PASS" if abs(diff) <= target_total * 0.1 else "REVISE"

    return report


def validate_citations(
    chapter_path: str,
    references_path: str | None = None,
) -> dict:
    """引用标记与文献列表一致性检查。

    Args:
        chapter_path: 章节 Markdown 文件路径。
        references_path: 参考文献列表文件路径（可选）。
    """
    try:
        path, text = safe_read(chapter_path)
    except ValueError as e:
        return {"error": str(e)}

    # 提取正文中的引用标记
    bracket_citations = set(re.findall(r'\[(\d+)\]', text))
    author_year = re.findall(r'[\(（]([^\)）]{2,40}(?:,|，)\s*\d{4})[\)）]', text)

    issues: list[str] = []
    cited_numbers = sorted(int(n) for n in bracket_citations if n.isdigit())

    if cited_numbers:
        max_num = max(cited_numbers)
        all_nums = set(range(1, max_num + 1))
        missing = all_nums - set(cited_numbers)
        if missing:
            issues.append(f"引用编号缺失: {sorted(missing)}")

    if references_path:
        try:
            ref_path, ref_text = safe_read(references_path)
        except ValueError as e:
            issues.append(f"参考文献文件错误: {e}")
        else:
            ref_entries = re.findall(r'^\s*[\[(]\d+[\])]', ref_text, re.MULTILINE)
            ref_nums: set[int] = set()
            for entry in ref_entries:
                nums = re.findall(r'\d+', entry)
                ref_nums.update(int(n) for n in nums if n.isdigit())

            in_text_not_in_ref = set(cited_numbers) - ref_nums
            if in_text_not_in_ref:
                issues.append(f"正文引用但参考文献列表缺失: {sorted(in_text_not_in_ref)}")

            in_ref_not_in_text = ref_nums - set(cited_numbers)
            if in_ref_not_in_text:
                issues.append(f"参考文献列表存在但正文未引用: {sorted(in_ref_not_in_text)}")

    return {
        "file": chapter_path,
        "cited_numbers": cited_numbers,
        "author_year_citations": author_year,
        "issues": issues,
        "status": "PASS" if not issues else "REVISE",
    }


def validate_chapter_citations(
    chapter_path: str,
    registry_path: str | None = None,
    cards_dir: str | None = None,
) -> dict:
    """章级引用编号与注册表一致性检查。

    针对 copilot 逐章撰写场景设计。优先按“当前章已更新注册表”检查；
    如注册表中没有“首次出现章节”信息，则退化为旧注册表差集模式。
    不检查 1..max 连续性（跨章编号不连续是正常的）。

    Args:
        chapter_path: 章节 Markdown 文件路径。
        registry_path: 引用编号注册表路径。
        cards_dir: 文献索引卡目录（可选，用于交叉验证 cite_key）。
    """
    try:
        path, text = safe_read(chapter_path)
    except ValueError as e:
        return {"error": str(e), "status": "REVISE"}

    cited_numbers = sorted({int(n) for n in re.findall(r'\[(\d+)\]', text) if n.isdigit()})
    if not cited_numbers:
        return {"file": chapter_path, "cited_numbers": [], "issues": [], "status": "PASS"}

    issues: list[str] = []

    if not registry_path:
        expected = list(range(1, len(cited_numbers) + 1))
        if cited_numbers != expected:
            issues.append(
                f"无注册表时，本章引用应从 [1] 开始连续递增；实际为 {cited_numbers}，期望 {expected}"
            )
        return {
            "file": chapter_path,
            "cited_numbers": cited_numbers,
            "issues": issues,
            "status": "PASS" if not issues else "REVISE",
        }

    try:
        _, reg_text = safe_read(registry_path)
    except ValueError as e:
        issues.append(f"注册表文件读取失败: {e}")
        return {"file": chapter_path, "cited_numbers": cited_numbers, "issues": issues, "status": "REVISE"}

    reg_entries = _parse_registry_entries(reg_text)
    if not reg_entries:
        return {
            "file": chapter_path,
            "cited_numbers": cited_numbers,
            "issues": ["注册表中未找到编号映射条目"],
            "status": "REVISE",
        }

    registry_nums: set[int] = set()
    registry_map: dict[int, str] = {}
    cite_key_to_num: dict[str, int] = {}
    duplicate_cite_keys: list[str] = []

    for entry in reg_entries:
        num = entry["num"]
        cite_key = entry["cite_key"]
        registry_nums.add(num)
        registry_map[num] = cite_key.strip()
        if cite_key.strip() in cite_key_to_num:
            duplicate_cite_keys.append(cite_key.strip())
        else:
            cite_key_to_num[cite_key.strip()] = num

    if duplicate_cite_keys:
        for ck in duplicate_cite_keys:
            issues.append(f"注册表内部不一致: cite_key '{ck}' 绑定了多个编号")

    cited_set = set(cited_numbers)
    unknown_nums = sorted(cited_set - registry_nums)
    if unknown_nums:
        issues.append(f"本章存在注册表中未登记的编号: {unknown_nums}")

    chapter_label = _extract_chapter_label(chapter_path, text)
    current_chapter_nums = {
        entry["num"]
        for entry in reg_entries
        if entry.get("first_chapter")
        and chapter_label
        and _normalize_chapter_label(entry["first_chapter"]) == chapter_label
    }
    validation_mode = "current_chapter_registry" if current_chapter_nums else "legacy_diff_fallback"
    extra_registry_nums = sorted(current_chapter_nums - cited_set)
    if extra_registry_nums:
        issues.append(
            f"注册表标记为本章首次出现的编号未在正文中引用: {extra_registry_nums}"
        )

    if current_chapter_nums:
        old_registry_nums = registry_nums - current_chapter_nums
        old_max = max(old_registry_nums) if old_registry_nums else 0
        expected_new = list(range(old_max + 1, old_max + 1 + len(current_chapter_nums)))
        current_sorted = sorted(current_chapter_nums)
        if current_sorted != expected_new:
            issues.append(
                f"本章新增编号不连续: 注册表中本章新增编号为 {current_sorted}，"
                f"但应从 {old_max + 1} 开始连续编号（期望 {expected_new}）"
            )
    else:
        # 兼容旧注册表模式：若注册表尚未更新，则只能从差集推断新增编号。
        registry_max = max(registry_nums) if registry_nums else 0
        new_in_chapter = cited_set - registry_nums
        if new_in_chapter:
            expected_new = list(range(registry_max + 1, registry_max + 1 + len(new_in_chapter)))
            new_sorted = sorted(new_in_chapter)
            if new_sorted != expected_new:
                issues.append(
                    f"新编号不连续: 本章引入的新编号为 {new_sorted}，"
                    f"但应从 {registry_max + 1} 开始连续编号（期望 {expected_new}）"
                )

    registry_max = max(registry_nums) if registry_nums else 0

    if cards_dir:
        try:
            from paper_tools_mcp.tools.literature import get_literature_cards
            cards_result = get_literature_cards(cards_dir)
            if "error" in cards_result:
                issues.append(f"索引卡交叉验证失败: {cards_result['error']}")
            elif "cards" in cards_result:
                card_cite_keys = {
                    c.get("cite_key", c.get("file", "")) for c in cards_result["cards"]
                }
                for num, ck in registry_map.items():
                    if ck not in card_cite_keys:
                        issues.append(
                            f"注册表编号 [{num}] 的 cite_key '{ck}' "
                            f"在索引卡目录中未找到对应卡片"
                        )
            else:
                issues.append("索引卡交叉验证失败: 返回结果缺少 cards 字段")
        except Exception as exc:
            issues.append(f"索引卡交叉验证失败: {exc}")

    return {
        "file": chapter_path,
        "cited_numbers": cited_numbers,
        "registry_max": registry_max,
        "chapter_label": chapter_label,
        "validation_mode": validation_mode,
        "current_chapter_nums": sorted(current_chapter_nums),
        "issues": issues,
        "status": "PASS" if not issues else "REVISE",
    }


def validate_citation_order(
    doc_path: str,
    registry_path: str | None = None,
    ref_heading: str | None = None,
) -> dict:
    """全文级引用排序检查。

    对含参考文献章节的完整文档，验证参考文献列表排序是否匹配
    文中引用的首次出现顺序，并做实体映射验证。

    Args:
        doc_path: 完整文档 Markdown 路径（含参考文献章节）。
        registry_path: 引用编号注册表路径（可选，用于实体映射验证）。
        ref_heading: 参考文献章节标题正则（可选，默认 `` 参考文[献献录] ``）。
    """
    try:
        _, text = safe_read(doc_path)
    except ValueError as e:
        return {"error": str(e), "status": "REVISE"}

    heading_pattern = ref_heading or r'参考文[献献录]'
    ref_section_match = re.search(
        rf'^#+\s*{heading_pattern}', text, re.MULTILINE,
    )

    if not ref_section_match:
        return {
            "file": doc_path,
            "status": "REVISE",
            "issues": [f"未找到参考文献章节（heading={heading_pattern}）"],
            "order_check": {"violations": []},
        }

    body_text = text[:ref_section_match.start()]
    ref_text = text[ref_section_match.start():]

    body_citations = re.findall(r'\[(\d+)\]', body_text)
    if not body_citations:
        return {
            "file": doc_path,
            "status": "PASS",
            "issues": [],
            "order_check": {"violations": []},
        }

    first_appearance: dict[int, int] = {}
    for m in re.finditer(r'\[(\d+)\]', body_text):
        num = int(m.group(1))
        if num not in first_appearance:
            first_appearance[num] = m.start()

    expected_order = sorted(first_appearance, key=first_appearance.get)

    ref_entries = re.findall(
        r'^\s*[\[(]\s*(\d+)\s*[\])]', ref_text, re.MULTILINE,
    )
    actual_order = [int(n) for n in ref_entries]

    violations: list[dict] = []
    issues: list[str] = []
    missing_in_list = [num for num in expected_order if num not in actual_order]
    extra_in_list = [num for num in actual_order if num not in expected_order]

    if missing_in_list:
        issues.append(f"参考文献列表缺失正文已引用条目: {missing_in_list}")
    if extra_in_list:
        issues.append(f"参考文献列表存在正文未引用条目: {extra_in_list}")

    for ref_num in sorted(set(expected_order) & set(actual_order)):
        expected_pos = expected_order.index(ref_num) + 1
        actual_pos = actual_order.index(ref_num) + 1
        if expected_pos != actual_pos:
            violations.append({
                "ref_num": ref_num,
                "expected_pos": expected_pos,
                "actual_pos": actual_pos,
                "description": (
                    f"参考文献 [{ref_num}] 位置错误：按首次出现顺序应在第 {expected_pos} 位，"
                    f"实际在第 {actual_pos} 位"
                ),
            })

    for v in violations:
        issues.append(v["description"])

    entity_issues: list[str] = []
    entity_check: dict = {}

    if registry_path:
        try:
            _, reg_text = safe_read(registry_path)
        except ValueError as exc:
            entity_issues.append(f"注册表读取失败: {exc}")
        else:
            reg_map: dict[int, str] = {}
            for entry in _parse_registry_entries(reg_text):
                num = entry["num"]
                ck = entry["cite_key"]
                if num not in reg_map:
                    reg_map[num] = ck.strip()

            try:
                from paper_tools_mcp.tools.literature import get_literature_cards
                card_dir = str(Path(registry_path).parent.parent / "04_参考文献" / "literature_cards")
                cards_result = get_literature_cards(card_dir)
                if "error" in cards_result:
                    entity_issues.append(f"实体映射检查失败: {cards_result['error']}")
                    cards_result = {"cards": []}
                ck_author_map: dict[str, str] = {}
                for c in cards_result.get("cards", []):
                    ck_author_map[c.get("cite_key", "")] = c.get("author", "")

                ref_author_map: dict[int, str] = {}
                for m in re.finditer(
                    r'^\s*[\[(]\s*(\d+)\s*[\])]\s*(.{2,30}?)[\.\,，。]', ref_text, re.MULTILINE,
                ):
                    num = int(m.group(1))
                    author_frag = m.group(2).strip()
                    ref_author_map[num] = author_frag

                for num, ck in reg_map.items():
                    reg_author = ck_author_map.get(ck, "")
                    ref_author = ref_author_map.get(num, "")
                    if reg_author and ref_author:
                        reg_author_surname = reg_author.split(",")[0].split("，")[0][:4]
                        ref_author_frag = ref_author[:4]
                        if reg_author_surname and ref_author_frag:
                            if not (
                                reg_author_surname.lower() in ref_author_frag.lower()
                                or ref_author_frag.lower() in reg_author_surname.lower()
                            ):
                                entity_issues.append(
                                    f"实体映射不一致: [{num}] 注册表 cite_key '{ck}' "
                                    f"对应作者 '{reg_author}'，但参考文献列表条目作者为 '{ref_author}'"
                                )
            except Exception as exc:
                entity_issues.append(f"实体映射检查失败: {exc}")

            entity_check = {
                "registry_entries": len(reg_map),
                "mismatches": entity_issues,
            }

    issues.extend(entity_issues)

    return {
        "file": doc_path,
        "status": "PASS" if not issues else "REVISE",
        "issues": issues,
        "order_check": {
            "expected_order": expected_order,
            "actual_order": actual_order,
            "violations": violations,
        },
        "entity_check": entity_check,
    }


def validate_format(
    chapter_path: str,
    spec_path: str,
) -> dict:
    """[DEPRECATED] 按 Markdown 格式规范检查。

    已拆分为 validate_markdown_structure + validate_assets。
    保留此函数仅为向后兼容，新代码请使用拆分后的函数。
    """
    result = validate_markdown_structure(chapter_path, spec_path)
    asset_result = validate_assets(chapter_path)
    result["assets_check"] = asset_result
    if asset_result.get("status") == "REVIEW":
        result["status"] = "REVIEW"
    return result


def validate_markdown_structure(
    chapter_path: str,
    spec_path: str,
) -> dict:
    """Markdown 结构合法性检查。

    检查标题层级、首行缩进、参考文献章节存在性、特殊标记完整性。
    不包含 DOCX 级检查（字体/行距/页边距由 validate_docx_* 系列处理）。

    Args:
        chapter_path: 章节 Markdown 文件路径。
        spec_path: 格式规范 Markdown 文件路径。
    """
    try:
        path, text = safe_read(chapter_path)
    except ValueError as e:
        return {"error": str(e)}

    try:
        _, spec_text = safe_read(spec_path)
    except ValueError as e:
        return {"error": str(e)}

    checks: list[dict] = []

    # 检查标题层级（检查是否跳级，如 # 直接到 ###）
    headings = re.findall(r'^(#{1,4})\s+(.+)$', text, re.MULTILINE)
    if headings:
        levels = [len(h[0]) for h in headings]
        skips = []
        for i in range(1, len(levels)):
            if levels[i] - levels[i - 1] > 1:
                skips.append(f"L{levels[i-1]} → L{levels[i]}（第 {i+1} 个标题）")
        checks.append({
            "rule": "标题层级",
            "status": "PASS" if not skips else "WARNING",
            "note": f"检测到 {len(headings)} 个标题" + (f"，跳级: {skips}" if skips else ""),
        })

    # 检查首行缩进
    paragraphs = [line for line in text.split('\n') if line.strip() and not line.startswith(('#', '|', '-', '*', '>', '`'))]
    indented = sum(1 for p in paragraphs if p.startswith((' ', '\t')))
    if paragraphs:
        indent_ratio = indented / len(paragraphs)
        checks.append({
            "rule": "首行缩进",
            "status": "PASS" if indent_ratio > 0.3 else "WARNING",
            "note": f"段落缩进比例: {indent_ratio:.0%}（{indented}/{len(paragraphs)}）",
        })

    # 检查参考文献章节存在性
    ref_section = re.search(r'#.*参考文[献献录]', text, re.IGNORECASE)
    if ref_section:
        ref_lines = text[ref_section.end():].strip().split('\n')
        ref_lines = [l for l in ref_lines if l.strip()]
        checks.append({
            "rule": "参考文献章节",
            "status": "PASS",
            "note": f"检测到 {len(ref_lines)} 条参考文献",
        })
    else:
        checks.append({
            "rule": "参考文献章节",
            "status": "WARNING",
            "note": "未找到参考文献章节（# 参考文献）",
        })

    # 检查特殊标记完整性（摘要、关键词、目录占位等）
    special_markers = {
        "摘要": r'#.*摘\s*要',
        "关键词": r'关键词[：:]',
    }
    for marker_name, pattern in special_markers.items():
        if re.search(pattern, text, re.IGNORECASE):
            checks.append({
                "rule": f"特殊标记: {marker_name}",
                "status": "PASS",
                "note": f"检测到{marker_name}标记",
            })
        else:
            # 不作为硬性失败，仅提示
            checks.append({
                "rule": f"特殊标记: {marker_name}",
                "status": "INFO",
                "note": f"未检测到{marker_name}标记",
            })

    issues = [c for c in checks if c["status"] in ("REVISE", "WARNING")]

    return {
        "file": chapter_path,
        "spec": spec_path,
        "checks": checks,
        "issues": issues,
        "status": "PASS" if not issues else "REVIEW",
    }


def validate_assets(
    chapter_path: str,
) -> dict:
    """Markdown 资源锚点检查。

    检查图片引用、表格语法、公式闭合性、脚注引用完整性。

    Args:
        chapter_path: 章节 Markdown 文件路径。
    """
    try:
        path, text = safe_read(chapter_path)
    except ValueError as e:
        return {"error": str(e)}

    file_dir = Path(chapter_path).resolve().parent
    assets: list[dict] = []
    broken: list[dict] = []

    # 检查图片引用
    images = re.findall(r'!\[([^\]]*)\]\(([^)]+)\)', text)
    for alt, img_path in images:
        full_path = (file_dir / img_path).resolve()
        assets.append({"type": "image", "alt": alt, "path": img_path})
        # 防止路径遍历：只检查沙箱内的文件
        try:
            full_path.relative_to(file_dir)
        except ValueError:
            broken.append({"type": "image", "alt": alt, "path": img_path, "issue": "路径在允许范围外"})
            continue
        if not full_path.exists():
            broken.append({"type": "image", "alt": alt, "path": img_path, "issue": "文件不存在"})

    # 检查表格语法（每行列数一致）
    table_blocks = re.findall(r'(\|.+\|[\r\n]+\|[-:| ]+[\r\n]+(?:\|.+\|[\r\n]*)+)', text)
    for idx, block in enumerate(table_blocks):
        rows = [line for line in block.strip().split('\n') if line.strip()]
        if rows:
            col_counts = [len([c for c in row.split('|') if c.strip() or True]) - 2 for row in rows]
            if len(set(col_counts)) > 1:
                broken.append({"type": "table", "index": idx + 1, "issue": f"列数不一致: {col_counts}"})
            assets.append({"type": "table", "index": idx + 1, "rows": len(rows), "cols": col_counts[0] if col_counts else 0})

    # 检查行内公式和块级公式闭合（先剥离代码块避免误报）
    # Step 1: 剥离三引号代码块
    stripped = re.sub(r'```[\s\S]*?```', '', text)
    # Step 2: 剥离行内代码
    stripped = re.sub(r'`[^`]+`', '', stripped)
    # Step 3: 剥离已转义的 \$
    stripped = re.sub(r'\\\$', '', stripped)

    # 统计 $ 数量（剥离后）
    dollar_count = stripped.count('$')
    # 跳过疑似货币符号（$ 后紧跟数字，前面不是字母）
    currency_pattern = re.findall(r'(?<![a-zA-Z])\$\d', stripped)
    effective_dollar_count = dollar_count - len(currency_pattern)

    if effective_dollar_count % 2 != 0:
        broken.append({"type": "formula", "issue": f"$ 符号数量为奇数（有效 {effective_dollar_count} 个），存在未闭合的公式"})

    # 行内公式计数（配对的 $...$）
    inline_math = re.findall(r'(?<!\$)\$(?!\$)([^\$]+?)\$', stripped)
    assets.append({"type": "inline_formula", "count": len(inline_math)})

    # 块级公式计数（配对的 $$...$$）
    block_math_pairs = len(re.findall(r'\$\$[^\$]*?\$\$', stripped))
    assets.append({"type": "block_formula", "count": block_math_pairs})

    # 检查脚注引用
    footnote_refs = re.findall(r'\[\^(\w+)\]', text)
    footnote_defs = re.findall(r'^\[\^(\w+)\]', text, re.MULTILINE)
    undefined_refs = set(footnote_refs) - set(footnote_defs)
    if undefined_refs:
        broken.append({"type": "footnote", "issue": f"未定义的脚注引用: {sorted(undefined_refs)}"})
    assets.append({"type": "footnote", "refs": len(footnote_refs), "defs": len(footnote_defs)})

    return {
        "file": chapter_path,
        "assets": assets,
        "broken": broken,
        "status": "PASS" if not broken else "REVIEW",
    }


# ── DOCX 级校验 ──────────────────────────────────────────────

# 标题样式匹配模式：同时覆盖 "Heading 1" 和中文数字 styleId
_HEADING_STYLE_PATTERNS = ["heading", "heading 1", "heading 2", "heading 3"]
_NUMERIC_STYLEID_HEADING_RE = re.compile(r"^[1-3]$")
_GENERIC_STYLES = {"Normal", "Header", "Footer", "No Spacing", "Default Paragraph Font"}
_LAYOUT_TOLERANCE_EMU = 12700  # 1pt 容差，应对 Word/OpenXML 舍入


def _load_docx(docx_path: str, template: bool = False) -> "Document | None":
    """安全加载 DOCX 文件，返回 Document 或 None。

    Args:
        docx_path: DOCX 文件路径。
        template: 是否为模板文件。模板路径走 template_path_exists，
            普通路径走 safe_path_exists。

    Returns:
        Document 对象，或 None（路径不存在 / python-docx 未安装 / 文件无效）。
    """
    path_check = template_path_exists if template else safe_path_exists
    if not path_check(docx_path):
        return None
    try:
        from docx import Document
        return Document(docx_path)
    except ImportError:
        return None
    except Exception:
        return None


def _is_heading_style(style_name: str, style_id: str | None) -> bool:
    """判断样式是否为标题样式（覆盖 'Heading 1' 和中文数字 styleId '1'/'2'/'3'）。"""
    if style_name and any(pattern in style_name.lower() for pattern in _HEADING_STYLE_PATTERNS):
        return True
    if style_id and _NUMERIC_STYLEID_HEADING_RE.match(style_id):
        return True
    return False


def validate_docx_styles(docx_path: str, spec_path: str) -> dict:
    """DOCX 样式命中检查（第一阶段实现）。

    检查：段落实际使用的样式分布、标题层级顺序、与模板的样式覆盖率。
    """
    doc = _load_docx(docx_path)
    if doc is None:
        return {
            "name": "validate_docx_styles",
            "status": "FAIL",
            "failure_codes": ["CLI_01"],
            "summary": f"无法打开 DOCX 文件: {docx_path}",
        }

    checks: list[dict] = []
    failure_codes: list[str] = []

    total_paras = len(doc.paragraphs)

    # 1. 段落样式分布统计
    paragraph_style_counts: dict[str, int] = {}
    heading_order: list[int] = []
    unstyled_count = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "None"
        paragraph_style_counts[style_name] = paragraph_style_counts.get(style_name, 0) + 1

        style_id = para.style.style_id if para.style else None
        if _is_heading_style(style_name, style_id):
            level = None
            if style_name and "heading" in style_name.lower():
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    pass
            if level is None and style_id:
                try:
                    level = int(style_id)
                except ValueError:
                    pass
            if level is not None:
                heading_order.append(level)
        elif style_name == "None":
            unstyled_count += 1

    # 2. 检查是否有标题段落（论文必须有标题结构）
    heading_count = sum(heading_order)
    if heading_count > 0:
        checks.append({
            "rule": "标题段落存在性",
            "status": "PASS",
            "detail": f"检测到 {heading_count} 个标题段落",
        })
    elif total_paras > 5:
        # 超过 5 段但无标题 → 排版可能失败
        checks.append({
            "rule": "标题段落存在性",
            "status": "FAIL",
            "detail": f"{total_paras} 段落中无标题段落",
        })
        failure_codes.append("DOCX_STYLE_01")

    # 3. 检查正文段落是否使用标准样式
    normal_count = paragraph_style_counts.get("Normal", 0)
    if total_paras > 0:
        normal_ratio = normal_count / total_paras
        checks.append({
            "rule": "正文段落样式命中",
            "status": "PASS" if normal_ratio > 0.3 else "REVIEW",
            "detail": f"Normal: {normal_count}/{total_paras} ({normal_ratio:.0%})",
        })
        if normal_ratio < 0.1 and total_paras > 5:
            failure_codes.append("DOCX_STYLE_02")

    # 4. 未使用样式的段落（排版失败的强信号）
    if unstyled_count > 0:
        checks.append({
            "rule": "无样式段落",
            "status": "REVIEW" if unstyled_count > total_paras * 0.1 else "INFO",
            "detail": f"{unstyled_count}/{total_paras} 段落未指定样式",
        })
        if unstyled_count > total_paras * 0.3:
            failure_codes.append("DOCX_STYLE_02")

    # 5. 检测标题层级是否合理（不应跳级）
    skip_found = False
    for i in range(1, len(heading_order)):
        if heading_order[i] - heading_order[i - 1] > 1:
            skip_found = True
            break
    if heading_order:
        checks.append({
            "rule": "标题层级顺序",
            "status": "PASS" if not skip_found else "REVIEW",
            "detail": f"标题级别序列: {heading_order[:10]}{'...' if len(heading_order) > 10 else ''}",
        })
        if skip_found:
            failure_codes.append("DOCX_STYLE_03")

    # 6. 与模板比对：检查模板独有样式是否在输出中被使用
    template_doc = _load_docx(spec_path, template=True) if spec_path else None
    if template_doc:
        template_style_names = {s.name for s in template_doc.styles if s.type == 1}  # PARAGRAPH only
        # 找模板中有但输出中段落未使用的样式
        used_output_styles = set(paragraph_style_counts.keys())
        missing_template_styles = template_style_names - used_output_styles
        # 过滤掉通用样式
        generic_styles = {"Normal", "Header", "Footer", "No Spacing"}
        missing_meaningful = [s for s in missing_template_styles if s not in generic_styles]
        if missing_meaningful:
            checks.append({
                "rule": "模板样式覆盖率",
                "status": "INFO",
                "detail": f"模板中有 {len(missing_meaningful)} 个段落样式未被输出使用: {missing_meaningful[:5]}",
            })

    return {
        "name": "validate_docx_styles",
        "status": "PASS" if not failure_codes else ("REVIEW" if all(c in ("DOCX_STYLE_02", "DOCX_STYLE_03") for c in failure_codes) else "FAIL"),
        "failure_codes": failure_codes,
        "summary": f"检查 {len(checks)} 项样式规则" + (f"，{len(failure_codes)} 个问题" if failure_codes else "，全部通过"),
        "checks": checks,
    }


def validate_docx_sections(docx_path: str, spec_path: str) -> dict:
    """DOCX 分节/分页检查（第一阶段实现）。

    检查页面尺寸、方向、页边距，与模板比对。
    """
    doc = _load_docx(docx_path)
    if doc is None:
        return {
            "name": "validate_docx_sections",
            "status": "FAIL",
            "failure_codes": ["CLI_01"],
            "summary": f"无法打开 DOCX 文件: {docx_path}",
        }

    checks: list[dict] = []
    failure_codes: list[str] = []

    # 加载模板做比对
    template_doc = None
    if spec_path:
        template_doc = _load_docx(spec_path, template=True)

    sections = doc.sections
    checks.append({"rule": "分节数量", "status": "PASS", "detail": f"{len(sections)} 个节"})

    if template_doc:
        template_sections = template_doc.sections
        if len(sections) != len(template_sections):
            checks.append({
                "rule": "分节数量一致性",
                "status": "REVIEW",
                "detail": f"输出 {len(sections)} 节 vs 模板 {len(template_sections)} 节",
            })
            failure_codes.append("DOCX_SECTION_01")

        # 逐节比对页面设置（不再只比较第一节）
        compare_count = min(len(sections), len(template_sections))
        margin_fields = [
            ("top_margin", "上边距"),
            ("bottom_margin", "下边距"),
            ("left_margin", "左边距"),
            ("right_margin", "右边距"),
        ]
        size_fields = [("page_width", "页面宽度"), ("page_height", "页面高度")]

        for idx in range(compare_count):
            out_sec = sections[idx]
            tpl_sec = template_sections[idx]
            section_label = f"section[{idx}]"

            # 页面方向
            if out_sec.orientation != tpl_sec.orientation:
                checks.append({
                    "rule": f"{section_label} 页面方向",
                    "status": "REVIEW",
                    "detail": f"输出 {out_sec.orientation} vs 模板 {tpl_sec.orientation}",
                })
                failure_codes.append("DOCX_SECTION_01")

            # 页边距（带容差）
            for field, label in margin_fields:
                out_val = getattr(out_sec, field)
                tpl_val = getattr(tpl_sec, field)
                diff_emu = abs(int(out_val) - int(tpl_val))
                if diff_emu > _LAYOUT_TOLERANCE_EMU:
                    checks.append({
                        "rule": f"{section_label} {label}",
                        "status": "REVIEW",
                        "detail": f"输出 {out_val / 914400:.2f}in vs 模板 {tpl_val / 914400:.2f}in（差 {diff_emu / 914400:.2f}in）",
                    })
                    failure_codes.append("DOCX_LAYOUT_01")

            # 页面尺寸（带容差）
            for field, label in size_fields:
                out_val = getattr(out_sec, field)
                tpl_val = getattr(tpl_sec, field)
                diff_emu = abs(int(out_val) - int(tpl_val))
                if diff_emu > _LAYOUT_TOLERANCE_EMU:
                    checks.append({
                        "rule": f"{section_label} {label}",
                        "status": "REVIEW",
                        "detail": f"输出 {out_val / 914400:.2f}in vs 模板 {tpl_val / 914400:.2f}in",
                    })
                    failure_codes.append("DOCX_LAYOUT_01")

        # 报告多余节
        if len(sections) > len(template_sections):
            for idx in range(len(template_sections), len(sections)):
                checks.append({
                    "rule": f"section[{idx}] 多余节",
                    "status": "REVIEW",
                    "detail": "输出中存在模板没有的节",
                })

    else:
        # 无模板时仅做基本报告
        for sec in sections:
            checks.append({
                "rule": "节属性",
                "status": "INFO",
                "detail": f"尺寸 {sec.page_width / 914400:.1f}x{sec.page_height / 914400:.1f}in, "
                         f"边距 上{sec.top_margin / 914400:.2f} 下{sec.bottom_margin / 914400:.2f} "
                         f"左{sec.left_margin / 914400:.2f} 右{sec.right_margin / 914400:.2f}in",
            })

    return {
        "name": "validate_docx_sections",
        "status": "PASS" if not failure_codes else "REVIEW",
        "failure_codes": failure_codes,
        "summary": f"检查 {len(sections)} 个节" + (f"，{len(failure_codes)} 个问题" if failure_codes else "，全部通过"),
        "checks": checks,
    }


def validate_docx_fonts(docx_path: str, spec_path: str) -> dict:
    """DOCX 字体命中检查。

    检查段落字体设置和 eastAsia 字体声明。
    当前仅执行基础检查（eastAsia 声明覆盖率）；
    spec_path 参数保留用于后续实现模板字体对比。
    """
    doc = _load_docx(docx_path)
    if doc is None:
        return {
            "name": "validate_docx_fonts",
            "status": "FAIL",
            "failure_codes": ["CLI_01"],
            "summary": f"无法打开 DOCX 文件: {docx_path}",
        }

    checks: list[dict] = []
    failure_codes: list[str] = []
    total_runs = 0
    missing_east_asia_count = 0

    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            total_runs += 1
            rpr = run._element.find(f"{{{DOCX_NS}}}rPr")
            rfonts = rpr.find(f"{{{DOCX_NS}}}rFonts") if rpr is not None else None
            east_asia = rfonts.get(f"{{{DOCX_NS}}}eastAsia") if rfonts is not None else None
            if not east_asia:
                missing_east_asia_count += 1

    if total_runs == 0:
        checks.append({"rule": "eastAsia 字体声明", "status": "INFO", "detail": "未检测到非空 run"})
    else:
        ratio = missing_east_asia_count / total_runs
        if ratio > 0.5:
            checks.append({
                "rule": "eastAsia 字体声明",
                "status": "REVIEW",
                "detail": f"{missing_east_asia_count}/{total_runs} 个非空 run 缺少 eastAsia 字体声明",
            })
            failure_codes.append("DOCX_FONT_01")
        else:
            checks.append({
                "rule": "eastAsia 字体声明",
                "status": "PASS",
                "detail": f"{missing_east_asia_count}/{total_runs} 个非空 run 缺少 eastAsia 声明（比例 {ratio:.0%}）",
            })

    return {
        "name": "validate_docx_fonts",
        "status": "REVIEW" if failure_codes else "PASS",
        "failure_codes": failure_codes,
        "summary": f"检查 {total_runs} 个非空 run 的字体声明" + (f"，{len(failure_codes)} 个问题" if failure_codes else "，通过"),
        "checks": checks,
    }


def validate_docx_references(docx_path: str) -> dict:
    """DOCX 参考文献段落样式检查。

    检查参考文献章节是否存在、段落是否使用一致样式。
    """
    doc = _load_docx(docx_path)
    if doc is None:
        return {
            "name": "validate_docx_references",
            "status": "FAIL",
            "failure_codes": ["CLI_01"],
            "summary": f"无法打开 DOCX 文件: {docx_path}",
        }

    checks: list[dict] = []
    failure_codes: list[str] = []

    # 查找参考文献章节
    ref_start = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.search(r'参考文[献献录]', text):
            ref_start = i
            break

    if ref_start is None:
        checks.append({
            "rule": "参考文献章节存在性",
            "status": "REVIEW",
            "detail": "未找到参考文献章节",
        })
        failure_codes.append("DOCX_REF_01")
        return {
            "name": "validate_docx_references",
            "status": "REVIEW",
            "failure_codes": failure_codes,
            "summary": "未找到参考文献章节",
            "checks": checks,
        }

    ref_paras = doc.paragraphs[ref_start + 1:]
    ref_count = len([p for p in ref_paras if p.text.strip()])

    if ref_count == 0:
        checks.append({"rule": "参考文献条目", "status": "REVIEW", "detail": "参考文献章节下无条目"})
        failure_codes.append("DOCX_REF_01")
    else:
        # 检查参考文献样式一致性
        ref_styles: dict[str, int] = {}
        for p in ref_paras:
            if not p.text.strip():
                continue
            sname = p.style.name if p.style else "None"
            ref_styles[sname] = ref_styles.get(sname, 0) + 1

        dominant = max(ref_styles, key=ref_styles.get) if ref_styles else "N/A"
        total_ref = sum(ref_styles.values())
        dominant_ratio = ref_styles[dominant] / total_ref if total_ref > 0 else 0

        checks.append({
            "rule": "参考文献条目",
            "status": "PASS",
            "detail": f"{ref_count} 条参考文献，主要样式: {dominant} ({dominant_ratio:.0%})",
        })
        if dominant_ratio < 0.7 and total_ref > 3:
            checks.append({
                "rule": "参考文献样式一致性",
                "status": "REVIEW",
                "detail": f"样式分布: {ref_styles}",
            })
            failure_codes.append("DOCX_REF_01")

    return {
        "name": "validate_docx_references",
        "status": "REVIEW" if failure_codes else "PASS",
        "failure_codes": failure_codes,
        "summary": f"检查 {ref_count} 条参考文献" + (f"，{len(failure_codes)} 个问题" if failure_codes else "，通过"),
        "checks": checks,
    }


def validate_docx_layout(docx_path: str, spec_path: str) -> dict:
    """DOCX 页面布局检查。

    检查所有节的页面尺寸是否一致（A4 vs Letter 等）。
    详细逐节边距检查由 validate_docx_sections 执行。
    spec_path 参数保留用于后续实现逐节边距模板比对。
    """
    doc = _load_docx(docx_path)
    if doc is None:
        return {
            "name": "validate_docx_layout",
            "status": "FAIL",
            "failure_codes": ["CLI_01"],
            "summary": f"无法打开 DOCX 文件: {docx_path}",
        }

    checks: list[dict] = []
    failure_codes: list[str] = []

    if not doc.sections:
        return {
            "name": "validate_docx_layout",
            "status": "FAIL",
            "failure_codes": ["DOCX_LAYOUT_01"],
            "summary": "文档无任何节",
            "checks": checks,
        }

    # 检查所有节页面尺寸一致性
    first_w = doc.sections[0].page_width
    first_h = doc.sections[0].page_height
    for idx, sec in enumerate(doc.sections):
        if sec.page_width != first_w or sec.page_height != first_h:
            checks.append({
                "rule": f"section[{idx}] 页面尺寸不一致",
                "status": "REVIEW",
                "detail": f"{sec.page_width / 914400:.1f}x{sec.page_height / 914400:.1f}in vs 第一节 {first_w / 914400:.1f}x{first_h / 914400:.1f}in",
            })
            failure_codes.append("DOCX_LAYOUT_01")

    if not failure_codes:
        checks.append({
            "rule": "页面尺寸一致性",
            "status": "PASS",
            "detail": f"全部 {len(doc.sections)} 节均为 {first_w / 914400:.1f}x{first_h / 914400:.1f}in",
        })

    return {
        "name": "validate_docx_layout",
        "status": "REVIEW" if failure_codes else "PASS",
        "failure_codes": failure_codes,
        "summary": f"检查 {len(doc.sections)} 个节的页面布局" + (f"，{len(failure_codes)} 个问题" if failure_codes else "，通过"),
        "checks": checks,
    }


def post_validate_docx(
    docx_path: str,
    spec_path: str,
    validation_mode: Literal["strict", "compat"] = "strict",
) -> dict:
    """汇总所有 docx 级校验结果（统一入口）。

    Args:
        docx_path: 输出 DOCX 文件路径。
        spec_path: 模板 DOCX 路径（用于 section/style 比对）。
        validation_mode: "strict"（REVIEW→FAILED）或 "compat"（REVIEW→REVIEW）。

    状态聚合规则：
        strict: FAIL→FAILED, REVIEW→FAILED, NOT_IMPLEMENTED→FAILED
        compat: FAIL→FAILED, REVIEW→REVIEW, NOT_IMPLEMENTED→REVIEW
    """
    checks = [
        validate_docx_styles(docx_path, spec_path),
        validate_docx_sections(docx_path, spec_path),
        validate_docx_fonts(docx_path, spec_path),
        validate_docx_references(docx_path),
        validate_docx_layout(docx_path, spec_path),
    ]

    overall_failure_codes: list[str] = []
    has_fail = False
    has_review_or_notimpl = False

    for check in checks:
        codes = check.get("failure_codes", [])
        overall_failure_codes.extend(codes)
        status = check.get("status", "UNKNOWN")
        if status == "FAIL":
            has_fail = True
        elif status in ("REVIEW", "NOT_IMPLEMENTED"):
            has_review_or_notimpl = True

    # 原始聚合
    if has_fail:
        raw_status = "FAIL"
    elif has_review_or_notimpl:
        raw_status = "REVIEW"
    else:
        raw_status = "PASS"

    # validation_mode 门控
    if validation_mode == "strict":
        final_status = "FAILED" if raw_status != "PASS" else "PASS"
    else:  # compat
        final_status = "FAILED" if raw_status == "FAIL" else ("REVIEW" if raw_status == "REVIEW" else "PASS")

    return {
        "docx_path": docx_path,
        "status": final_status,
        "raw_status": raw_status,
        "validation_mode": validation_mode,
        "checks": checks,
        "overall_failure_codes": overall_failure_codes,
    }
